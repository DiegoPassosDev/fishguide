"""Convert FishGuide .docx specification documents to Markdown — v2."""

from docx import Document
from docx.oxml.ns import qn
import os, re

DOCS_DIR = 'C:\\Users\\dspxr\\OneDrive\\Documentos\\Projetos\\React\\FishGuide\\documentos'
OUT_DIR = 'C:\\Users\\dspxr\\OneDrive\\Documentos\\Projetos\\React\\FishGuide\\docs'


def is_heading_text(text: str) -> bool:
    t = text.strip()
    if re.match(r'^\d+(\.\d+)*\.?\s+\w', t):
        return True
    if re.match(r'^[A-ZÀ-Ú][A-ZÀ-Ú\s\.]{4,}$', t) and len(t) > 8:
        return True
    known = ['Objetivo', 'Introdução', 'Conclusão', 'Resumo', 'Escopo',
             'Filosofia', 'Princípios', 'Visão Geral', 'Manifesto']
    if any(t.startswith(kw) for kw in known):
        return True
    return False


def heading_level(text: str) -> int:
    t = text.strip()
    if re.match(r'^\d+\.\d+\.\d+\s', t): return 3
    if re.match(r'^\d+\.\d+\s', t): return 2
    if re.match(r'^\d+\.\s', t): return 2
    if t.isupper() and len(t) > 10: return 1
    return 2


def is_meta(text: str) -> bool:
    t = text.lower().strip()
    prefixes = ['projeto', 'versão', 'versao', 'autor', 'data', 'status']
    return any(t.startswith(p + ':') or t.startswith(p + ' ') for p in prefixes)


def collect_paragraph_text(para) -> str:
    """Build paragraph text from runs, merging adjacent bold/plain correctly."""
    parts = []
    # Group consecutive runs with same bold state
    runs = para.runs
    if not runs:
        return para.text

    i = 0
    while i < len(runs):
        run = runs[i]
        text = run.text
        if not text:
            i += 1
            continue

        is_bold = run.bold
        # Look ahead to see how many consecutive runs share the bold state
        j = i + 1
        while j < len(runs):
            if runs[j].bold == is_bold or not runs[j].text.strip():
                j += 1
            else:
                break

        combined = ''.join(r.text for r in runs[i:j])

        if is_bold:
            # Check if this is a heading (all bold paragraph)
            # We'll only wrap in ** if it's NOT the whole paragraph being a heading
            parts.append(f'**{combined}**')
        else:
            parts.append(combined)

        i = j

    return ''.join(parts)


def is_entirely_bold(para) -> bool:
    """Check if entire paragraph content is bold."""
    texts = [r.text for r in para.runs if r.text.strip()]
    bolds = [r.bold for r in para.runs if r.text.strip()]
    if not texts:
        return False
    return all(bolds)


def table_to_md(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ''

    col_widths = []
    for ci in range(len(rows[0])):
        col_widths.append(max(len(r[ci]) for r in rows) + 2)
    lines = []
    header = '| ' + ' | '.join(r.center(w) for r, w in zip(rows[0], col_widths)) + ' |'
    lines.append(header)
    sep = '|-' + '-|-'.join('-' * w for w in col_widths) + '-|'
    lines.append(sep)
    for row in rows[1:]:
        line = '| ' + ' | '.join(r.ljust(w) for r, w in zip(row, col_widths)) + ' |'
        lines.append(line)
    return '\n'.join(lines)


def convert_docx_to_md(filepath: str) -> str:
    doc = Document(filepath)

    lines = []
    title = ''
    meta_lines = []
    in_list = False
    prev_text = ''

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            if in_list:
                lines.append('')
                in_list = False
            prev_text = ''
            continue

        full_text = collect_paragraph_text(para)
        entirely_bold = is_entirely_bold(para)

        if i == 0:
            title = text
            lines.append(f'# {title}')
            lines.append('')
            continue

        if is_meta(text):
            meta_lines.append(text)
            continue

        if meta_lines and (is_heading_text(text) or i > 15):
            for m in meta_lines:
                key, _, val = m.partition(':')
                if not val:
                    key, _, val = m.partition(' ')
                key = key.strip()
                val = val.strip()
                if key and val:
                    lines.append(f'- **{key.capitalize()}:** {val}')
                else:
                    lines.append(m)
            lines.append('')
            meta_lines = []

        if entirely_bold and is_heading_text(text):
            in_list = False
            level = heading_level(text)
            lines.append(f'{"#" * level} {text}')
            lines.append('')
            prev_text = text
            continue

        # List detection
        t_clean = text.lstrip('-*').strip()
        is_list = False
        if prev_text.endswith(':') or prev_text.endswith(';'):
            is_list = True
        elif t_clean and t_clean[0].islower() and len(t_clean) < 60:
            is_list = True
        elif text.startswith('- ') or text.startswith('* '):
            is_list = True

        if is_list:
            if not in_list:
                in_list = True
            lines.append(f'- {full_text}')
            prev_text = text
            continue
        else:
            if in_list:
                in_list = False
                lines.append('')

        lines.append(full_text)
        lines.append('')
        prev_text = text

    # Remaining metadata
    if meta_lines:
        for m in meta_lines:
            lines.append(m)
        lines.append('')

    # Tables at the end
    for table in doc.tables:
        lines.append(table_to_md(table))
        lines.append('')

    return '\n'.join(lines).strip()


def slugify(num, title=''):
    return f'{int(num):02d}'


def main():
    files = sorted([f for f in os.listdir(DOCS_DIR) if f.endswith('.docx') and not f.startswith('~')])
    converted = []
    errors = []

    for fname in files:
        src = os.path.join(DOCS_DIR, fname)
        print(f'Converting: {fname}...', end=' ')
        try:
            md = convert_docx_to_md(src)
            num = re.search(r'(\d+)', fname)
            n = num.group(1) if num else '99'
            out_name = f'{slugify(n)}.md'
            dst = os.path.join(OUT_DIR, out_name)
            with open(dst, 'w', encoding='utf-8') as f:
                f.write(md)
            converted.append((out_name, len(md)))
            print(f'OK ({len(md)} chars)')
        except Exception as e:
            errors.append((fname, str(e)))
            print(f'ERROR: {e}')

    print(f'\n=== Summary ===')
    print(f'Converted: {len(converted)} files')
    if errors:
        print(f'Errors: {len(errors)}')
        for n, e in errors:
            print(f'  {n}: {e}')
    print(f'Output: {OUT_DIR}')


if __name__ == '__main__':
    main()
